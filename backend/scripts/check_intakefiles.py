"""A client's file, before it is ever a client's file.

Two halves, in the order they were built. The first is the assertion that had to
land *before* an anonymous caller could send a file at all: `app/attachments.py`
opens a `.docx` and an `.xlsx` with a zip reader, and a zip reader with no bound
on what it unpacks will happily turn a quarter of a megabyte on the wire into
two hundred off it.

The second half, appended below the first, is `app/intakefiles.py` itself -
where an intake's files live, how a caller-supplied id is gated before it can
reach a path or a bucket key, what the client's own caps are, and what `close()`
takes with it. It runs against the **local** backend, because that is what
`intakefiles.configured() == False` gives you on a machine with no Spaces
credentials, and that fallback existing is the point. The Spaces branch is
exercised too, against a stub client rather than the network: what is proven
there is that the branch is reached and what it would send, not that
DigitalOcean would accept it.

That was tolerable while every file reaching `attachments.read` came from a
signed-in studio member uploading their own tender pack. It stops being
tolerable the moment the uploader is a stranger holding a link, which is what
the rest of this plan builds.

The bound is a sum of the sizes the archive itself declares, read from the
central directory, before any reader opens the archive. That is sound rather
than merely convenient, and the section at the bottom of this file proves the
premise it rests on rather than asserting it: CPython's `zipfile` truncates
every entry's read at the size that entry declares (`ZipExtFile._read1` does
`data = data[:self._left]`), so a declared size is an upper bound on what any
reader built on `zipfile` - python-docx and openpyxl both are - can obtain. An
archive that understates its own sizes does not get to expand past them; it
fails its CRC instead.

**Nothing here may raise into a request.** `attachments.py`'s module docstring
is explicit that a file which cannot be read is reported on the quotation as a
file that could not be read, and a client who is not in the room to be asked
makes that rule stronger, not weaker. So every refusal below is checked as an
`Attachment` carrying a `problem`, never as an exception.

    cd backend
    .venv/Scripts/python.exe scripts/check_intakefiles.py
"""

from __future__ import annotations

import io
import os
import struct
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
os.environ["GENERATED_DIR"] = tempfile.mkdtemp(prefix="prism-intakefiles-")
# Off, because every check in this project runs OFFLINE. The brief check is a
# real Gemini call on the generation path; left on, these scripts would reach
# the network, cost money, and fail on a machine with no key. `app/main.py`
# reads the flag at call time, so this line is the whole of the opt-out.
os.environ["CHECK_BRIEF_IS_REAL"] = "0"
# Blanked for the same reason every other check script blanks them: `app.config`
# reads these once at import time via `load_dotenv(..., override=False)`, and
# this repo's real `backend/.env` names an actual Supabase project. Nothing in
# this file sends a request today, but Task 2's sections will.
os.environ["SUPABASE_URL"] = ""
os.environ["SUPABASE_ANON_KEY"] = ""
os.environ["SUPABASE_JWT_SECRET"] = ""
# And all five Spaces settings, for a sharper reason than the three above, and
# one that is no longer hypothetical: `backend/.env` on this machine holds real
# DigitalOcean credentials for a real bucket. The local backend is only under
# test at all while `intakefiles.configured()` is False, so without these four
# lines half of this file would silently retarget at that bucket - and the half
# in question saves objects and then deletes prefixes. Blanked here, before
# `app.config` reads the environment, so what runs is what this file says it
# runs. Do not remove them to "test against the real thing"; write a separate
# script for that, under a prefix of its own.
os.environ["DO_SPACES_ACCESS_KEY"] = ""
os.environ["DO_SPACES_SECRET_KEY"] = ""
os.environ["DO_SPACES_REGION"] = ""
os.environ["DO_SPACES_BUCKET"] = ""
os.environ["DO_SPACES_ENDPOINT"] = ""

import docx  # noqa: E402
import openpyxl  # noqa: E402

from app import attachments  # noqa: E402

FAILURES: list[str] = []


def ok(label: str, condition: bool) -> None:
    print(("ok    " if condition else "FAIL  ") + label)
    if not condition:
        FAILURES.append(label)


def declared_total(data: bytes) -> int:
    """What the archive says it unpacks to, summed over every entry."""
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return sum(entry.file_size for entry in archive.infolist())


# --- The fixtures ------------------------------------------------------------
#
# Every bomb below is a *real* office document with a real payload bolted on,
# not a hand-forged archive with lying metadata. That matters: refusing a file
# whose central directory merely claims to be enormous would prove only that
# the metadata is read. These fixtures genuinely unpack to what they declare,
# so a bound that is not there really does hand two hundred megabytes to a
# reader. They are built by streaming the payload through `ZipFile.open(...,
# "w")` a megabyte at a time rather than materialising it - the whole point is
# that a bomb is cheap to make and expensive to open, and this script should
# only pay the cheap half.

BOMB_MEGABYTES = 192
ONE_MEGABYTE = b"\0" * (1024 * 1024)


def _with_bomb(base: bytes, name: str) -> bytes:
    out = io.BytesIO(base)
    with zipfile.ZipFile(out, "a", zipfile.ZIP_DEFLATED) as archive:
        with archive.open(name, "w") as handle:
            for _ in range(BOMB_MEGABYTES):
                handle.write(ONE_MEGABYTE)
    return out.getvalue()


def plain_docx() -> bytes:
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("A scope of work, in the shape a client actually sends one.")
    document.save(buffer)
    return buffer.getvalue()


def plain_xlsx() -> bytes:
    buffer = io.BytesIO()
    book = openpyxl.Workbook()
    book.active.append(["Item", "Quantity"])
    book.active.append(["Booking module", 1])
    book.save(buffer)
    return buffer.getvalue()


PLAIN_DOCX = plain_docx()
PLAIN_XLSX = plain_xlsx()
BOMB_DOCX = _with_bomb(PLAIN_DOCX, "word/media/payload.bin")
BOMB_XLSX = _with_bomb(PLAIN_XLSX, "xl/media/payload.bin")

ok(
    "the docx bomb is small on the wire and enormous off it - the shape of the "
    "attack, not a large upload",
    len(BOMB_DOCX) < 1_000_000 and declared_total(BOMB_DOCX) > 190 * 1024 * 1024,
)
ok(
    "the xlsx bomb is the same shape",
    len(BOMB_XLSX) < 1_000_000 and declared_total(BOMB_XLSX) > 190 * 1024 * 1024,
)


# --- The controls: an ordinary document still reads --------------------------
#
# First, because a bound that refuses everything would pass every refusal test
# in this file and break the feature it is protecting.

read_docx = attachments.read("scope.docx", PLAIN_DOCX)
ok(
    "an ordinary .docx still reads, with no problem reported",
    read_docx.problem == "" and "shape a client actually sends one" in read_docx.text,
)

read_xlsx = attachments.read("items.xlsx", PLAIN_XLSX)
ok(
    "an ordinary .xlsx still reads, with no problem reported",
    read_xlsx.problem == "" and "Booking module" in read_xlsx.text,
)


# --- The bomb is refused, and reported rather than raised --------------------

bombed_docx = attachments.read("scope.docx", BOMB_DOCX)
ok(
    "a .docx that unpacks past the bound is refused",
    bombed_docx.problem != "",
)
ok(
    "and refused the way this module refuses everything - an Attachment with no "
    "text and a reason, never an exception",
    bombed_docx.text == "" and not bombed_docx.usable,
)
ok(
    "and the reason names the file, the way every other problem in this module does",
    bombed_docx.problem.startswith("scope.docx"),
)

bombed_xlsx = attachments.read("items.xlsx", BOMB_XLSX)
ok(
    "an .xlsx that unpacks past the bound is refused too - the two readers are "
    "different libraries and the bound has to sit ahead of both",
    bombed_xlsx.problem != "" and bombed_xlsx.text == "",
)


# --- Refused *rather than read*, which is the whole point --------------------
#
# The assertions above would pass if the bomb were opened, expanded to two
# hundred megabytes, and only then reported. What this task is closing is the
# expansion itself, so the claim has to be that neither reader was ever
# reached. Both `_from_docx` and `_from_xlsx` do their `import` inside the
# function body and call through the module (`docx.Document(...)`,
# `openpyxl.load_workbook(...)`), so patching the attribute on the library
# module is visible to the running reader.

_real_document = docx.Document
_real_load_workbook = openpyxl.load_workbook
opens = {"docx": 0, "xlsx": 0}


def _counting_document(*args, **kwargs):
    opens["docx"] += 1
    return _real_document(*args, **kwargs)


def _counting_load_workbook(*args, **kwargs):
    opens["xlsx"] += 1
    return _real_load_workbook(*args, **kwargs)


try:
    docx.Document = _counting_document
    openpyxl.load_workbook = _counting_load_workbook
    attachments.read("scope.docx", BOMB_DOCX)
    attachments.read("items.xlsx", BOMB_XLSX)
    ok(
        "the docx bomb never reached docx.Document at all - the bound sits ahead "
        "of the reader, not after it",
        opens["docx"] == 0,
    )
    ok(
        "and the xlsx bomb never reached openpyxl.load_workbook",
        opens["xlsx"] == 0,
    )

    # The same instrument, pointed the other way: an ordinary document must
    # still get all the way to its reader, or the two assertions above would
    # pass on a module that had simply stopped reading anything.
    attachments.read("scope.docx", PLAIN_DOCX)
    attachments.read("items.xlsx", PLAIN_XLSX)
    ok(
        "an ordinary .docx does still reach docx.Document - the counters above "
        "are measuring a refusal, not a module that reads nothing",
        opens["docx"] == 1,
    )
    ok("and an ordinary .xlsx does still reach openpyxl.load_workbook", opens["xlsx"] == 1)
finally:
    docx.Document = _real_document
    openpyxl.load_workbook = _real_load_workbook


# --- A file that is not an archive keeps the message it already had ----------
#
# The bound must not swallow the ordinary corrupt-file case. A `.docx` that is
# not a zip at all cannot be inspected for declared sizes, and the right answer
# is to let the reader fail and report what it has always reported - not to
# invent a second message for a file whose problem is that it is broken, not
# that it is large.

not_an_archive = attachments.read("scope.docx", b"This is plainly not a zip archive.")
ok(
    "a .docx that is not an archive still gets this module's existing wording, "
    "not a new one about size",
    not_an_archive.problem == "scope.docx could not be opened.",
)

truncated = attachments.read("items.xlsx", BOMB_XLSX[: len(BOMB_XLSX) // 2])
ok(
    "and a half-written archive is still just a file that could not be opened",
    truncated.problem == "items.xlsx could not be opened.",
)


# --- The bound is a bound, at its exact edge ---------------------------------
#
# Proven by moving the bound rather than by building a fixture of exactly the
# right size: the constant is read at call time, so patching it is the precise
# instrument. `PLAIN_DOCX` unpacks to a few tens of kilobytes; set the bound to
# exactly that and it must still read, set it one byte lower and it must not.

ok(
    "attachments names a bound of its own for what an archive may unpack to",
    hasattr(attachments, "MAX_UNPACKED_BYTES"),
)

if hasattr(attachments, "MAX_UNPACKED_BYTES"):
    _real_bound = attachments.MAX_UNPACKED_BYTES
    plain_total = declared_total(PLAIN_DOCX)
    try:
        attachments.MAX_UNPACKED_BYTES = plain_total
        ok(
            "an archive that unpacks to exactly the bound is read - the bound is a "
            "ceiling, not a fence one byte below it",
            attachments.read("scope.docx", PLAIN_DOCX).problem == "",
        )
        attachments.MAX_UNPACKED_BYTES = plain_total - 1
        ok(
            "and one byte past it is refused",
            attachments.read("scope.docx", PLAIN_DOCX).problem != "",
        )
    finally:
        attachments.MAX_UNPACKED_BYTES = _real_bound

    ok(
        "the real bound admits an ordinary document with room to spare",
        attachments.MAX_UNPACKED_BYTES > 16 * 1024 * 1024,
    )
    ok(
        "and refuses the bombs above",
        attachments.MAX_UNPACKED_BYTES < declared_total(BOMB_DOCX),
    )


# --- The premise the bound rests on, proven rather than assumed --------------
#
# Summing what the central directory declares is only a real bound if a reader
# cannot get more out of an entry than that entry declared. CPython's
# `ZipExtFile._read1` truncates each decompressed block to the number of bytes
# still owed against the declared size (`data = data[:self._left]`), so it
# cannot. Asserted here, against the running interpreter, so that a future
# CPython which changed that would fail this script rather than quietly leave
# `attachments.py` with a decorative bound.
#
# The fixture is a genuine fifty-megabyte entry whose declared size has been
# rewritten to a hundred bytes - the only shape that could defeat a
# metadata-based bound, if the metadata were advisory.

_honest = io.BytesIO()
with zipfile.ZipFile(_honest, "w", zipfile.ZIP_DEFLATED) as archive:
    archive.writestr("word/document.xml", b"A" * (50 * 1024 * 1024))
_liar = _honest.getvalue().replace(
    struct.pack("<I", 50 * 1024 * 1024), struct.pack("<I", 100)
)

with zipfile.ZipFile(io.BytesIO(_liar)) as archive:
    ok(
        "the lying fixture really does understate itself",
        archive.infolist()[0].file_size == 100,
    )
    try:
        recovered = archive.read("word/document.xml")
        understated_result = f"returned {len(recovered)} bytes"
    except zipfile.BadZipFile:
        understated_result = "raised BadZipFile"

ok(
    "an archive that understates its own sizes cannot be read past them - it "
    "fails its CRC instead, so the declared sum really is an upper bound",
    understated_result == "raised BadZipFile",
)


# =============================================================================
# Task 2: where a client's file lives
# =============================================================================
#
# Everything below is `app/intakefiles.py`. It is appended here rather than
# woven into the sections above because those sections patch `docx.Document`
# and `openpyxl.load_workbook` on the library modules themselves and assert
# exact call counts - anything inserted between that block's `try:` and its
# `finally:` that happened to read a document would break assertions with
# nothing to do with it.

from app import config, intakefiles, intakes, storage, workspaces  # noqa: E402

ok(
    "importing app.intakefiles imports neither boto3 nor botocore - every check "
    "in this directory runs offline with no credentials, and a module-level "
    "import would make intakes.py, and so the whole app, unimportable on a "
    "machine that has not installed them yet. Both, not just boto3: the retry "
    "and timeout bound is a botocore.config.Config, and it is just as absent",
    "boto3" not in sys.modules and "botocore" not in sys.modules,
)


def refuses(label: str, call) -> None:
    """The assertion for a call that must be turned away, not merely survived."""
    try:
        call()
    except (intakefiles.IntakeFileError, intakes.IntakeError):
        ok(label, True)
        return
    ok(label, False)


workspaces.ensure_ready()
WORKSPACE = workspaces.create("Neptune Labs")
workspaces.use(WORKSPACE.id)


def new_intake(scope: str) -> intakes.Intake:
    return intakes.create(
        client_email="buyer@client.com",
        client_phone="",
        scope=scope,
        budget_text="",
        preset={},
        created_by="riku@neptune.ph",
    )


#: Not real files. Nothing below decodes them - `attachments.py` is what reads a
#: client's bytes and it is tested against real documents above. What this half
#: of the file is about is where the bytes go and who may address them, and for
#: that a recognisable byte string is a better fixture than a real PNG, because
#: an assertion that the bytes came back unchanged is legible on sight.
PNG = b"\x89PNG\r\n\x1a\n" + b"pretend this is a photograph of a site"
PDF = b"%PDF-1.4 pretend this is a scope of work"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# --- Unconfigured is the whole test surface, and that is deliberate ----------

ok(
    "with no credentials on the machine, Spaces is not configured - so the "
    "local backend is what answers, which is exactly the fallback the plan "
    "exists to make testable before credentials do",
    intakefiles.configured() is False,
)


# --- What save() hands back, which is the manifest entry itself --------------

FIRST = new_intake("A booking site for two clinics.")
SECOND = new_intake("A different client entirely.")

#: Every entry saved against FIRST, in the order they were saved. The listing
#: and deletion assertions below compare against this rather than against a
#: hand-counted number: a section added in the middle of this file should not
#: silently make "listing() reports one row per file" assert the wrong count, and
#: when it does drift it should be obvious which side is wrong.
FIRST_FILES: list[dict] = []


def save_to_first(name: str, data: bytes, kind: str, **extra) -> dict:
    entry = intakefiles.save(FIRST.id, name, data, kind, **extra)
    FIRST_FILES.append(entry)
    return entry


shot = save_to_first("site photo.png", PNG, "image/png")

ok(
    "save() returns exactly the five keys the record holds and no sixth",
    set(shot) == {"id", "name", "kind", "bytes", "note"},
)
ok(
    "the file is addressed by a server-minted 12-hex id, the same shape "
    "storage.is_valid_id already gates every other id in this codebase on",
    isinstance(shot["id"], str) and storage.is_valid_id(shot["id"]),
)
ok(
    "the client's own filename is kept, as data",
    shot["name"] == "site photo.png",
)
ok(
    "kind is the content type, canonicalised - one string that is the manifest's, "
    "the stored object's and read()'s all at once",
    shot["kind"] == "image/png",
)
ok(
    "bytes is an int, and is what actually arrived",
    isinstance(shot["bytes"], int) and shot["bytes"] == len(PNG),
)
ok(
    "note is empty when the file read cleanly",
    shot["note"] == "",
)

# --- A long name keeps the one part of it anything downstream reads ----------
#
# `clean_name` caps at 120 characters and the extension is at the far end of
# that string - which is the end a naive `[:120]` removes. Everything that
# decides what a file *is* keys on it: `resolve_type` falls back to the suffix
# when a browser declared nothing useful, and `attachments.read` picks its
# reader from the suffix and nothing else. A 200-character `.docx` used to be
# stored as `<id>.bin` under `application/octet-stream` and shown to the studio
# as a blob. Asserted through `save()` rather than against `clean_name` alone,
# because the bug was in what `save()` then resolved off the shortened string.
long_named = save_to_first(
    "Requirements - " + "a" * 200 + ".docx",
    PLAIN_DOCX,
    "application/octet-stream",
)
ok(
    "a 200-character filename is capped at 120 and still ends in its extension",
    len(long_named["name"]) == 120 and long_named["name"].endswith(".docx"),
)
ok(
    "so a browser that declared nothing useful still gets the type the extension "
    "names, not application/octet-stream",
    long_named["kind"]
    == "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
ok(
    "and a long name with no extension at all is simply cut, not given one",
    intakefiles.clean_name("b" * 200) == "b" * 120,
)
ok(
    "a dot in the middle of a long name is not treated as an extension - keeping "
    "it would spend the whole budget on the wrong end of the string",
    intakefiles.clean_name("c" * 100 + "." + "d" * 100) == ("c" * 100 + "." + "d" * 19),
)
ok(
    "and the kept extension is spelled the way it arrived - a name is data, and "
    "a truncation should not case-fold what a shorter name keeps",
    intakefiles.clean_name("E" * 200 + ".PDF").endswith(".PDF"),
)
ok(
    "and a note travels when the caller has one - it is where attachments.read's "
    "warning about a scan with no text layer ends up",
    save_to_first("scan.pdf", PDF, "application/pdf", note="no text layer")["note"]
    == "no text layer",
)
ok(
    "no URL and no bucket key anywhere in the entry - both are derived from the "
    "ids at read time, so a presigned URL cannot go stale on the record and a "
    "backend swap does not rewrite history",
    not any(
        "http" in str(value)
        or f"{intakefiles.KEY_PREFIX}/" in str(value)
        or intakefiles.DIRNAME in str(value)
        or FIRST.id in str(value)
        for value in shot.values()
    ),
)


# --- Resolving what a file actually is ---------------------------------------
#
# The declared type wins when it says something. `application/octet-stream` does
# not say anything - it is in `CONTENT_TYPES` only so `.bin` has a type to map
# back to - and treating it as an answer was a real defect rather than a
# hypothetical one. It is reachable through Task 3: `_read_documents` gates on
# the filename suffix and never inspects `content_type`, so any browser that
# declares octet-stream for a `.docx` (several do) would have had it stored as
# `<id>.bin` with `kind: application/octet-stream`, and Task 5 would show the
# studio a generic blob where a Word document was sent. Text extraction survived
# only because `attachments.read` dispatches on the name - luck, not design.

octet_docx = save_to_first("scope.docx", PLAIN_DOCX, "application/octet-stream")
ok(
    "a .docx that a browser declared as application/octet-stream is stored as a "
    ".docx, not as a nameless blob",
    octet_docx["kind"] == DOCX_MIME,
)
ok(
    "and comes back that way, so the studio is shown a Word document",
    intakefiles.read(FIRST.id, octet_docx["id"])[1] == DOCX_MIME,
)
ok(
    "a file with no recognisable type at all is still stored, as octet-stream - "
    "refusing a type is the upload route's job, where there is a person to tell",
    save_to_first("mystery.zzz", PDF, "")["kind"] == "application/octet-stream",
)

# The other half of that fallback, which is an asymmetry on purpose. A suffix is
# chosen by whoever uploaded the file, so it may earn a *document* type but never
# a raster one - `inline` is the only thing in this module a client could
# otherwise talk their way into, and `payload.png` should not be the argument
# that wins it. Task 3 gates images on the declared type, so nothing legitimate
# takes this path; the assertion is here so that a later task changing that gate
# finds this decision already made rather than making it by accident.
ok(
    "a suffix alone cannot earn a raster type, and so cannot earn inline: an "
    "executable named payload.png is stored as octet-stream and as an attachment",
    save_to_first("payload.png", PDF, "application/x-msdownload")["kind"]
    == "application/octet-stream",
)
ok(
    "while a genuinely declared image still gets one",
    intakefiles.disposition_for("image/png") == "inline"
    and intakefiles.disposition_for("application/octet-stream") == "attachment"
    and intakefiles.disposition_for("image/svg+xml") == "attachment",
)


# --- The client's filename is a string, never a path -------------------------

sneaky = save_to_first("../../../../etc/passwd", PDF, "application/pdf")
ok(
    "a filename that is a traversal is kept as a name and never used as one",
    sneaky["name"] == "passwd" and storage.is_valid_id(sneaky["id"]),
)
windows = save_to_first(r"C:\Users\someone\scope.docx", PDF, "application/pdf")
ok(
    "and a drive-absolute one is the same story",
    windows["name"] == "scope.docx",
)

FILES_DIR = workspaces.root() / intakefiles.DIRNAME / FIRST.id
ok(
    "the bytes land under _intake_files/<intake_id>/, beside _intakes/",
    FILES_DIR.is_dir(),
)
ok(
    "every file on disk is named <file_id>.<ext> - nothing of what the client "
    "called it reaches a path",
    all(storage.is_valid_id(path.stem) for path in FILES_DIR.iterdir()),
)
ok(
    "and intakes.listing() still cannot see the directory: it globs *.json in "
    "_intakes/, so a sibling folder is invisible to the intake walk",
    len(intakes.listing()) == 2,
)


# --- Reading it back, and the round trip that has to hold --------------------

recovered = intakefiles.read(FIRST.id, shot["id"])
ok(
    "read() gives the bytes back unchanged",
    recovered is not None and recovered[0] == PNG,
)
ok(
    "and the content type it reports is exactly the kind the manifest recorded - "
    "one string in three places, so nothing downstream has to guess which is right",
    recovered is not None and recovered[1] == shot["kind"],
)
pdf_back = intakefiles.read(FIRST.id, sneaky["id"])
ok(
    "the same round trip for a document",
    pdf_back == (PDF, "application/pdf"),
)

rows = intakefiles.listing(FIRST.id)
ok(
    "listing() reports one row per file actually stored, and one per file saved",
    {row["id"] for row in rows} == {entry["id"] for entry in FIRST_FILES}
    and len(rows) == len(FIRST_FILES),
)
ok(
    "as {id, kind, bytes} and not the manifest's five - the client's filename "
    "and the extraction note live on the record, because storage never had them",
    all(set(row) == {"id", "kind", "bytes"} for row in rows),
)
ok(
    "with the size and the type storage itself holds, matching what save() "
    "recorded for each - the third place that one canonical type has to agree",
    {row["id"]: (row["kind"], row["bytes"]) for row in rows}
    == {entry["id"]: (entry["kind"], entry["bytes"]) for entry in FIRST_FILES},
)

ok(
    "unconfigured, view_url is the app's own route rather than a presigned one - "
    "the frontend links to one thing and never learns which backend answered",
    intakefiles.view_url(FIRST.id, shot["id"]) == f"/api/intakes/{FIRST.id}/files/{shot['id']}",
)


# --- The id gate, in both positions ------------------------------------------
#
# The load-bearing section. Each of these must come back as "no" without ever
# having built a path or a key, and the last of them - a file id that belongs to
# a different intake - is the one the plan names as carrying the whole feature.

BAD_INTAKE_IDS = [
    "../" + FIRST.id,
    "..",
    "../../secret",
    "..\\..\\secret",
    # A drive-absolute segment. `Path("x") / "C:/Windows/..."` discards the left
    # side entirely on Windows, so ungated this reaches a real system path and
    # the call raises `PermissionError` rather than answering anything - which
    # is why the loop below catches: a mutation of the gate should print a clean
    # FAIL, not take the whole report down with a traceback three sections
    # before the end.
    "C:/Windows/System32/config/sam",
    "/etc/passwd",
    "",
    "   ",
    "not-hex-at-all",
    FIRST.id + ".json",
    FIRST.id[:11],
    FIRST.id + "0",
]


def answers_nothing(intake_id: str) -> bool:
    """Every read-side function turns this id away, and none of them raises."""
    try:
        return (
            intakefiles.read(intake_id, shot["id"]) is None
            and intakefiles.listing(intake_id) == []
            and intakefiles.view_url(intake_id, shot["id"]) == ""
            and intakefiles.forget(intake_id) == 0
        )
    except Exception as exc:  # noqa: BLE001 - a raise here is a failure, not a crash
        print(f"      ({intake_id!r} raised {exc.__class__.__name__}: {exc})")
        return False


ok(
    "an intake id that is not a bare 12-hex string - a traversal, a drive-absolute "
    "segment, an empty string, a near miss - reads nothing, lists nothing, mints "
    "no URL and deletes nothing",
    all(answers_nothing(bad) for bad in BAD_INTAKE_IDS),
)
refuses(
    "and save() refuses it outright rather than inventing somewhere to put the bytes",
    lambda: intakefiles.save("../../secret", "scope.pdf", PDF, "application/pdf"),
)

# Everything above is a weak assertion and it is worth saying so rather than
# letting it read as proof. Eleven of those twelve ids resolve to a directory
# that does not exist, so "answers nothing" holds with no gate at all - remove
# `storage.is_valid_id` from `_intake_key` and they all still pass. They are kept
# because they are cheap and they document the shape being refused; they are not
# what proves the gate.
#
# This is. `_intakes/` is one sibling up from `_intake_files/`, it really does
# exist, and it holds one `<intake_id>.json` per record. An intake id is itself a
# perfectly valid 12-hex *file* id, so the lookup's glob matches it. And an
# intake record carries `token` - the client's live bearer credential, the one
# field in this codebase deliberately excluded from every response.
#
# Ungated, measured: `read` returns that record with the token in it; `listing`
# returns every intake in the workspace with its size; and `forget` unlinks every
# intake record and rmdirs the directory - every client's token, every state,
# every quotation pointer, gone. That last one is the worst thing this gate
# prevents, and until now nothing asserted it.

ESCAPE = f"../{intakes.DIRNAME}"
records_before = {entry.id for entry in intakes.listing()}

ok(
    "a traversal that lands somewhere real - _intakes/, where every client's live "
    "link token is stored - reads nothing",
    intakefiles.read(ESCAPE, FIRST.id) is None,
)
ok(
    "and lists nothing, rather than enumerating every intake in the workspace",
    intakefiles.listing(ESCAPE) == [],
)
ok(
    "and mints no URL to one",
    intakefiles.view_url(ESCAPE, FIRST.id) == "",
)
ok(
    "and deletes nothing - ungated, this call unlinks every intake record in the "
    "workspace and removes the directory with them",
    intakefiles.forget(ESCAPE) == 0,
)
ok(
    "with every record still on disk after all four, which is the assertion that "
    "actually names the damage",
    {entry.id for entry in intakes.listing()} == records_before and len(records_before) >= 2,
)

BAD_FILE_IDS = ["../../../etc/passwd", "..", "", "   ", "not-hex", shot["id"][:11], "*"]
file_gate_holds = True
for bad in BAD_FILE_IDS:
    file_gate_holds = file_gate_holds and intakefiles.read(FIRST.id, bad) is None
    file_gate_holds = file_gate_holds and intakefiles.view_url(FIRST.id, bad) == ""
ok(
    "the same gate on the file id, which is the other half of the pair and would "
    "otherwise be a glob or a traversal inside a directory that is legitimately open",
    file_gate_holds,
)

ok(
    "a file id belonging to another intake returns None - the pair is checked "
    "together, because a real id in the wrong intake is the attack a valid-looking "
    "id was always going to be",
    intakefiles.read(SECOND.id, shot["id"]) is None,
)
ok(
    "and mints no URL for it either, which matters more: a presigned URL is a "
    "bearer credential and handing one out is the disclosure, whatever the route "
    "returns next",
    intakefiles.view_url(SECOND.id, shot["id"]) == "",
)
ok(
    "and the same instrument pointed the other way - the file is still readable "
    "from its own intake, so everything above is measuring a refusal rather than "
    "a module that has stopped reading anything",
    intakefiles.read(FIRST.id, shot["id"]) == (PNG, "image/png"),
)


# --- The caps, and that they are the client's rather than the studio's --------

ok(
    "config names a file count cap for a client, and it is tighter than what a "
    "studio may attach to its own pad",
    config.MAX_CLIENT_FILES < config.MAX_IMAGES + config.MAX_DOCUMENTS,
)
ok(
    "and a per-file byte cap strictly below both of the studio's own, so a studio "
    "raising one of its limits cannot raise a stranger's by doing it",
    config.MAX_CLIENT_FILE_BYTES < config.MAX_IMAGE_BYTES
    and config.MAX_CLIENT_FILE_BYTES < config.MAX_DOCUMENT_BYTES,
)
ok(
    "and there are exactly three MAX_CLIENT_* names, so no second aggregate cap "
    "was invented beside the one Task 1 already added - enumerated rather than "
    "inferred from an inequality that would have held either way",
    {name for name in vars(config) if name.startswith("MAX_CLIENT_")}
    == {"MAX_CLIENT_FILES", "MAX_CLIENT_FILE_BYTES", "MAX_CLIENT_UPLOAD_TOTAL_BYTES"},
)


# --- The manifest on the record ----------------------------------------------

THIRD = new_intake("The one the manifest is written onto.")
ok("an intake starts with nothing attached", THIRD.attachments == [])
ok(
    "attachments is a field advance() is allowed to write - /submit writes it "
    "through _client_advance on the one move it is legal from",
    "attachments" in intakes.ADVANCE_FIELDS,
)

manifest = [shot, sneaky]
intakes.advance(THIRD.id, intakes.SUBMITTED, attachments=manifest)
ok(
    "and the manifest survives the round trip through disk with every key and "
    "every type intact",
    intakes.get(THIRD.id).attachments == manifest,
)

FOURTH = new_intake("The one whose manifest is refused.")
refuses(
    "advance() refuses a manifest that is not a list at all - Intake.model_validate "
    "runs on the whole record, so a malformed write fails before anything is saved",
    lambda: intakes.advance(FOURTH.id, intakes.SUBMITTED, attachments="scope.pdf"),
)
refuses(
    "and one whose entries are not dicts",
    lambda: intakes.advance(FOURTH.id, intakes.SUBMITTED, attachments=["scope.pdf"]),
)
ok(
    "and the intake those two were refused against is untouched - still issued, "
    "with nothing attached. Asserted here, between the refusals and the write "
    "below, because after that write it would only be measuring the write",
    intakes.get(FOURTH.id).state == intakes.ISSUED
    and intakes.get(FOURTH.id).attachments == [],
)
ok(
    "but List[dict] is the whole of what the model enforces: the five keys are a "
    "convention held by intakefiles.save() being their only producer, not "
    "something the record checks, and this assertion is here to say so out loud",
    intakes.advance(FOURTH.id, intakes.SUBMITTED, attachments=[{"unrelated": 1}]).attachments
    == [{"unrelated": 1}],
)


# --- Deletion, through both doors to the same state ---------------------------

FIFTH = new_intake("Closed by close().")
kept = intakefiles.save(FIFTH.id, "scope.pdf", PDF, "application/pdf")
ok("a file is there to lose", intakefiles.read(FIFTH.id, kept["id"]) is not None)
intakes.close(FIFTH.id, "riku@neptune.ph")
ok(
    "close() takes the client's files with it - a withdrawn request keeping a "
    "stranger's scope document on disk is the other half of the decision that "
    "already blanks its token",
    intakefiles.read(FIFTH.id, kept["id"]) is None,
)
ok(
    "and leaves no directory behind",
    not (workspaces.root() / intakefiles.DIRNAME / FIFTH.id).exists(),
)

SIXTH = new_intake("Closed by advance().")
also = intakefiles.save(SIXTH.id, "scope.pdf", PDF, "application/pdf")
intakes.advance(SIXTH.id, intakes.CLOSED)
ok(
    "advance(id, CLOSED) is the other, equally legal door to that state and "
    "deletes them too - the deletion sits in _write's chokepoint beside the token "
    "blank, because two places implementing one decision is how they drift apart",
    intakefiles.read(SIXTH.id, also["id"]) is None,
)

ok(
    "forget() counts what it removed",
    intakefiles.forget(FIRST.id) == len(FIRST_FILES),
)
ok(
    "and a second call over the same intake is a no-op, not a failure",
    intakefiles.forget(FIRST.id) == 0,
)
ok(
    "with the directory gone with them",
    not FILES_DIR.exists(),
)


# --- The Spaces branch, against a stub rather than the network ---------------
#
# There are no Spaces credentials on this machine and this script does not want
# any: it must stay offline. What can still be proven, and is worth proving, is
# that the branch is reachable, that the key it builds is the one the plan
# specifies, and that the object it would write is private and carries the
# headers a later task will serve without being in the path itself. What is
# **not** proven here is that DigitalOcean would accept any of it.


class StubSpaces:
    """As much of the S3 client as the Spaces branch actually calls.

    It pages, and it can refuse individual keys on a delete. Both of those exist
    because the first version of this stub did neither, and so three code paths
    in `intakefiles.py` - the continuation loop, the `_DELETE_BATCH` flush and
    `_delete_batch`'s per-key error accounting - were in the diff and executed by
    nothing. `MAX_CLIENT_FILES` is 6, so none of the three is reachable in
    production; they are still exactly the kind of place an off-by-one or a
    `KeyError` waits, and untested code in a diff is untested code.
    """

    #: Small enough that three objects page. The real service returns up to a
    #: thousand and the number is not what is being tested - the loop is.
    PAGE = 2

    def __init__(self) -> None:
        self.objects: dict[str, dict] = {}
        self.puts: list[dict] = []
        self.presigned: list[tuple] = []
        self.pages_served = 0
        self.delete_calls: list[list[str]] = []
        self.fail_delete = False
        #: Keys this Space will report as undeletable, the way a real one does:
        #: a 200 whose body carries an `Errors` list, not an exception.
        self.undeletable: set[str] = set()

    def put_object(self, **kwargs):
        self.puts.append(dict(kwargs))
        self.objects[kwargs["Key"]] = dict(kwargs)
        return {}

    def list_objects_v2(self, **kwargs):
        prefix = kwargs.get("Prefix", "")
        after = kwargs.get("ContinuationToken", "")
        matching = sorted(key for key in self.objects if key.startswith(prefix))
        remaining = [key for key in matching if key > after] if after else matching
        page = remaining[: self.PAGE]
        truncated = len(remaining) > len(page)
        self.pages_served += 1
        answer = {
            "Contents": [{"Key": key, "Size": len(self.objects[key]["Body"])} for key in page],
            "KeyCount": len(page),
            "IsTruncated": truncated,
        }
        if truncated:
            answer["NextContinuationToken"] = page[-1]
        return answer

    def get_object(self, **kwargs):
        entry = self.objects.get(kwargs["Key"])
        if entry is None:
            raise LookupError("NoSuchKey")
        return {"Body": io.BytesIO(entry["Body"]), "ContentType": entry.get("ContentType", "")}

    def delete_objects(self, **kwargs):
        if self.fail_delete:
            raise ConnectionError("Spaces could not be reached.")
        keys = [item["Key"] for item in kwargs["Delete"]["Objects"]]
        self.delete_calls.append(list(keys))
        deleted, errors = [], []
        for key in keys:
            if key in self.undeletable:
                errors.append({"Key": key, "Code": "AccessDenied", "Message": "no"})
                continue
            self.objects.pop(key, None)
            deleted.append({"Key": key})
        return {"Deleted": deleted, "Errors": errors} if errors else {"Deleted": deleted}

    def generate_presigned_url(self, operation, Params=None, ExpiresIn=None):  # noqa: N803
        self.presigned.append((operation, dict(Params or {}), ExpiresIn))
        return f"https://prism-test.sgp1.digitaloceanspaces.com/{Params['Key']}?X-Amz-Expires={ExpiresIn}"


def unreachable_client():
    raise ConnectionError("boto3 could not build a client.")


STUB = StubSpaces()
_real_client = intakefiles._client
_real_credentials = (
    config.DO_SPACES_ACCESS_KEY,
    config.DO_SPACES_SECRET_KEY,
    config.DO_SPACES_REGION,
    config.DO_SPACES_BUCKET,
    config.DO_SPACES_ENDPOINT,
)

try:
    config.DO_SPACES_ACCESS_KEY = "stub-access-key"
    config.DO_SPACES_SECRET_KEY = "stub-secret-never-logged"
    config.DO_SPACES_REGION = "sgp1"
    config.DO_SPACES_BUCKET = "prism-test"
    config.DO_SPACES_ENDPOINT = ""
    intakefiles._client = lambda: STUB

    ok(
        "with the four that matter set, configured() is True - read live off "
        "config on every call, exactly as mailer.configured() is, so nothing "
        "caches a boolean taken at import time",
        intakefiles.configured() is True,
    )
    ok(
        "and the endpoint is not one of the four: a studio who set the credentials, "
        "the region and the bucket has configured Spaces whether or not they also "
        "wrote down a hostname this module can derive",
        intakefiles.configured() is True and config.DO_SPACES_ENDPOINT == "",
    )
    ok(
        "with nothing configured, the endpoint is derived from the region, so the "
        "ordinary case has nothing to get wrong",
        intakefiles.endpoint() == "https://sgp1.digitaloceanspaces.com",
    )
    config.DO_SPACES_ENDPOINT = "https://sgp1.cdn.digitaloceanspaces.com/"
    ok(
        "and a configured one wins, trailing slash trimmed - for the cases the "
        "derivation cannot reach: a CDN hostname, a region newer than that line, "
        "or an S3-compatible store that is not DigitalOcean",
        intakefiles.endpoint() == "https://sgp1.cdn.digitaloceanspaces.com",
    )
    config.DO_SPACES_ENDPOINT = ""

    ok(
        "the client is built with an explicit connect timeout, read timeout and "
        "retry bound rather than botocore's 60s/60s/5-attempt defaults - forget() "
        "runs inline on an async close route, so an unreachable bucket would park "
        "the event loop for roughly ten minutes across its two round trips, and "
        "every anonymous /submit with it",
        intakefiles.SPACES_CONNECT_TIMEOUT_SECONDS <= 10
        and intakefiles.SPACES_READ_TIMEOUT_SECONDS <= 30
        and intakefiles.SPACES_MAX_ATTEMPTS <= 3,
    )

    SEVENTH = new_intake("Filed in Spaces.")
    spaces_shot = intakefiles.save(SEVENTH.id, "site photo.png", PNG, "image/png")
    spaces_scope = intakefiles.save(SEVENTH.id, "scope of work.docx", PLAIN_DOCX, DOCX_MIME)

    written = STUB.puts[0]
    ok(
        "objects are written private, said explicitly rather than inherited from "
        "whatever the bucket happens to default to",
        written["ACL"] == "private",
    )
    ok(
        "with the validated content type on the object itself",
        written["ContentType"] == "image/png",
    )
    ok(
        "a raster image from the allowlist is the one thing marked inline - the "
        "studio opens these on the studio's own origin",
        written["ContentDisposition"] == "inline",
    )
    ok(
        "and everything outside that allowlist is written as an attachment, at "
        "put_object time, where a later task cannot forget it on read",
        STUB.puts[1]["ContentDisposition"] == "attachment",
    )
    ok(
        "no-store rides on the object too, so a presigned GET carries it without "
        "this app being in the path",
        written["CacheControl"] == "no-store",
    )
    # Built from `KEY_PREFIX` rather than spelling the first segment out, so
    # renaming the prefix is one edit rather than two - the literal version of
    # this line is what failed when it moved from `intakes` to `PRISM`, which is
    # the assertion doing its job but for the wrong reason. The SHAPE is what
    # this is about: app, then workspace, then intake, then a 12-hex file id
    # with the stored extension and nothing of the client's own filename.
    ok(
        "the key is <prefix>/<workspace>/<intake>/<file_id>.<ext>",
        written["Key"]
        == f"{intakefiles.KEY_PREFIX}/{WORKSPACE.id}/{SEVENTH.id}/{spaces_shot['id']}.png",
    )
    ok(
        "and the prefix really is the app's own name, so PRISM's objects are "
        "findable as a unit in a bucket that holds other applications' too",
        intakefiles.KEY_PREFIX == "PRISM",
    )
    ok(
        "and nothing of the client's own filename is in it",
        "site photo" not in written["Key"] and "scope of work" not in STUB.puts[1]["Key"],
    )
    ok(
        "no credential is written onto the object, and none is in the manifest entry",
        not any(
            isinstance(value, str) and ("stub-secret" in value or "stub-access-key" in value)
            for value in written.values()
        )
        and "stub-secret" not in str(spaces_shot),
    )
    ok(
        "nothing was written to local disk while Spaces was configured - the two "
        "backends are alternatives, not layers",
        not (workspaces.root() / intakefiles.DIRNAME / SEVENTH.id).exists(),
    )

    ok(
        "read() on the Spaces branch gives the bytes and the stored type back",
        intakefiles.read(SEVENTH.id, spaces_shot["id"]) == (PNG, "image/png"),
    )
    ok(
        "and a file id from another intake is unreadable here too - a bucket key "
        "is not a path and `..` does not traverse it, but the id gate is what "
        "stops one intake addressing another's objects, which is the same attack "
        "by a different road",
        intakefiles.read(FIRST.id, spaces_shot["id"]) is None,
    )
    ok(
        "listing() reads the prefix, and recovers each file id back out of a key "
        "that is `/`-separated whatever platform this is running on",
        [row["id"] for row in intakefiles.listing(SEVENTH.id)]
        == sorted([spaces_shot["id"], spaces_scope["id"]]),
    )
    ok(
        "with each one's stored type and size",
        {row["id"]: (row["kind"], row["bytes"]) for row in intakefiles.listing(SEVENTH.id)}
        == {
            spaces_shot["id"]: ("image/png", len(PNG)),
            spaces_scope["id"]: (DOCX_MIME, len(PLAIN_DOCX)),
        },
    )

    presigned = intakefiles.view_url(SEVENTH.id, spaces_shot["id"])
    operation, params, ttl = STUB.presigned[-1]
    ok(
        "configured, view_url is a presigned GET on that one key",
        presigned.startswith("https://")
        and operation == "get_object"
        and params["Key"] == written["Key"]
        and params["Bucket"] == "prism-test",
    )
    ok(
        "with a life measured in minutes, not hours",
        isinstance(ttl, int) and 0 < ttl <= 900,
    )
    ok(
        "and no URL is minted for a file id from another intake, which is checked "
        "before anything is signed rather than after",
        intakefiles.view_url(FIRST.id, spaces_shot["id"]) == "" and len(STUB.presigned) == 1,
    )

    EIGHTH = new_intake("Deleted from Spaces.")
    intakefiles.save(EIGHTH.id, "scope.pdf", PDF, "application/pdf")
    ok(
        "forget() deletes the whole prefix and says how many",
        intakefiles.forget(SEVENTH.id) == 2,
    )

    # --- The three paths that were in the diff and executed by nothing --------
    #
    # `MAX_CLIENT_FILES` is 6 and a page is a thousand, so none of these is
    # reachable in production. They are still code, and code nothing runs is
    # where an off-by-one waits. The instrument is the stub's own page size and
    # the module's `_DELETE_BATCH`, both moved rather than fixtures of a
    # thousand files built to reach them.

    PAGED = new_intake("Enough files to page.")
    paged_ids = sorted(
        intakefiles.save(PAGED.id, f"file{index}.pdf", PDF, "application/pdf")["id"]
        for index in range(5)
    )
    STUB.pages_served = 0
    ok(
        "listing() follows the continuation token across pages rather than "
        "returning only the first - five objects, two per page",
        [row["id"] for row in intakefiles.listing(PAGED.id)] == paged_ids
        and STUB.pages_served == 3,
    )
    ok(
        "and a single-object lookup still resolves through the same paged walk",
        intakefiles.read(PAGED.id, paged_ids[0]) == (PDF, "application/pdf"),
    )

    _real_batch = intakefiles._DELETE_BATCH
    try:
        intakefiles._DELETE_BATCH = 2
        STUB.delete_calls = []
        removed = intakefiles.forget(PAGED.id)
        ok(
            "forget() flushes a full batch mid-walk instead of accumulating every "
            "key - five objects at two per call is three calls, and every one of "
            "them is deleted",
            removed == 5
            and [len(call) for call in STUB.delete_calls] == [2, 2, 1]
            and intakefiles.listing(PAGED.id) == [],
        )
    finally:
        intakefiles._DELETE_BATCH = _real_batch

    REFUSED = new_intake("One key the Space will not give up.")
    stubborn = intakefiles.save(REFUSED.id, "stuck.pdf", PDF, "application/pdf")
    fine = intakefiles.save(REFUSED.id, "fine.pdf", PDF, "application/pdf")
    STUB.undeletable = {
        key for key in STUB.objects if key.endswith(f"{stubborn['id']}.pdf")
    }
    ok(
        "a Space that refuses one key reports it in the response body rather than "
        "raising, and forget() counts what actually went - one of two, not two, "
        "and not zero",
        intakefiles.forget(REFUSED.id) == 1,
    )
    ok(
        "with the key it could not remove still there, and the one it could gone",
        [row["id"] for row in intakefiles.listing(REFUSED.id)] == [stubborn["id"]]
        and fine["id"] not in {row["id"] for row in intakefiles.listing(REFUSED.id)},
    )
    STUB.undeletable = set()

    STUB.fail_delete = True
    ok(
        "a delete that cannot reach the network is logged and counted as nothing "
        "removed, never raised",
        intakefiles.forget(EIGHTH.id) == 0,
    )
    ok(
        "and close() still closes - the intake closing is the thing the studio "
        "asked for, and an unreachable bucket is not their problem to be told about",
        intakes.close(EIGHTH.id, "riku@neptune.ph").state == intakes.CLOSED
        and intakes.get(EIGHTH.id).state == intakes.CLOSED,
    )

    NINTH = new_intake("Closed with no client at all.")
    intakefiles.save(NINTH.id, "scope.pdf", PDF, "application/pdf")
    intakefiles._client = unreachable_client
    ok(
        "forget() survives a client that cannot even be built - the failure one "
        "step earlier than a failed delete, and the one a machine with bad "
        "credentials actually hits",
        intakefiles.forget(NINTH.id) == 0,
    )
    ok(
        "and that close() too",
        intakes.close(NINTH.id, "riku@neptune.ph").state == intakes.CLOSED,
    )
finally:
    intakefiles._client = _real_client
    (
        config.DO_SPACES_ACCESS_KEY,
        config.DO_SPACES_SECRET_KEY,
        config.DO_SPACES_REGION,
        config.DO_SPACES_BUCKET,
        config.DO_SPACES_ENDPOINT,
    ) = _real_credentials

ok(
    "configured() is False again the moment the credentials go away - live, not "
    "latched",
    intakefiles.configured() is False,
)
ok(
    "and neither boto3 nor botocore was imported, through any of that: the whole "
    "Spaces branch - put, get, list, presign, paginate, batch-delete, and every "
    "failure mode above - was exercised through the one seam that would have "
    "imported them",
    "boto3" not in sys.modules and "botocore" not in sys.modules,
)


print()
print(f"{len(FAILURES)} FAILED" if FAILURES else "all pass")
sys.exit(1 if FAILURES else 0)
